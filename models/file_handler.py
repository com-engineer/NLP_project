import os
from io import BytesIO
from docx import Document
import PyPDF2
import pdfplumber

# File processing libraries
try:
    import PyPDF2
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    print("PyPDF2 not available")
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    print("python-docx not available")
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    print("pdfplumber not available")
    PDFPLUMBER_AVAILABLE = False


class ResumeFileHandler:
    """Handles text extraction from various resume file formats"""

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']

    def extract_text_from_file(self, file_path=None, file_content=None, file_extension=None):
        """Extract text from various resume file formats"""

        if file_path:
            file_extension = os.path.splitext(file_path)[1].lower()
            with open(file_path, 'rb') as file:
                file_content = file.read()
        elif file_content and file_extension:
            file_extension = file_extension.lower()
        else:
            raise ValueError("Either file_path or (file_content + file_extension) must be provided")

        # Route to appropriate extraction method
        if file_extension == '.pdf':
            return self._extract_from_pdf(file_content)
        elif file_extension == '.docx':
            return self._extract_from_docx(file_content)
        elif file_extension == '.txt':
            return self._extract_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def _extract_from_pdf(self, file_content):
        """Extract text from PDF files"""
        text = ""

        if PDF_AVAILABLE:
            try:
                pdf_file = BytesIO(file_content)
                reader = PdfReader(pdf_file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if text.strip():
                    return text
            except Exception as e:
                print(f"PyPDF2 extraction failed: {e}")

        if PDFPLUMBER_AVAILABLE:
            try:
                pdf_file = BytesIO(file_content)
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    return text
            except Exception as e:
                print(f"pdfplumber extraction failed: {e}")

        if not text.strip():
            raise ValueError("Could not extract text from PDF file")
        return text

    def _extract_from_docx(self, file_content):
        """Extract text from DOCX files"""
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx library not available")

        try:
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)

            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return text
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX file: {e}")

    def _extract_from_txt(self, file_content):
        """Extract text from TXT files"""
        try:
            encodings = ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue
            text = file_content.decode('utf-8', errors='ignore')
            return text
        except Exception as e:
            raise ValueError(f"Failed to extract text from TXT file: {e}")

    def is_supported_format(self, file_extension):
        """Check if file format is supported"""
        return file_extension.lower() in self.supported_formats
